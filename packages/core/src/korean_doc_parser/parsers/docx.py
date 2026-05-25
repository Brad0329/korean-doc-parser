"""DOCX parser backed by python-docx (Phase 0 decision: self-implement).

Outputs:
* Markdown built from paragraphs + Heading-style → ``#`` mapping
* Tables (one ``ParsedTable`` per ``doc.tables`` entry)
* Images extracted by direct ZIP access to ``word/media/*`` —
  ``ExtractedImage.file_path`` points at a tempfile, ``sha256`` is computed.
* ``ParseMetadata`` from ``core_properties`` (title, author, created, modified)

DOCX has no native page concept, so :attr:`ParsedTable.page_no` and
:attr:`ExtractedImage.page_no` are left ``None``.

The parser registers itself with the global registry on import.
"""

from __future__ import annotations

import hashlib
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING

from korean_doc_parser.core import (
    BaseParser,
    ExtractedImage,
    ParsedTable,
    ParseMetadata,
    ParseResult,
    register_parser,
)
from korean_doc_parser.exceptions import ParseError

if TYPE_CHECKING:
    from docx.document import Document as _Document

__all__ = ["DocxParser"]


# Heading style names that should become Markdown headings.
# Both English (python-docx default) and Korean (한글 워드 사용자 환경) covered.
_HEADING_LEVELS: dict[str, int] = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
    "제목 1": 1,
    "제목 2": 2,
    "제목 3": 3,
    "제목 4": 4,
    "제목 5": 5,
    "제목 6": 6,
}


class DocxParser(BaseParser):
    """Parse ``.docx`` files using python-docx."""

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, path: Path) -> ParseResult:
        from docx import Document

        try:
            doc = Document(str(path))
            markdown = self._build_markdown(doc)
            tables = self._extract_tables(doc)
            images = self._extract_images(path)
            metadata = self._extract_metadata(doc, path)
            return ParseResult(
                markdown=markdown,
                metadata=metadata,
                tables=tables,
                images=images,
            )
        except ParseError:
            raise
        except Exception as exc:
            msg = f"Failed to parse DOCX {path}: {exc}"
            raise ParseError(msg) from exc

    # ─────────────────────────────────────────────────────────────────────
    # Helpers
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _build_markdown(doc: _Document) -> str:
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style is not None else ""
            level = _HEADING_LEVELS.get(style_name)
            if level is not None:
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        return "\n\n".join(parts)

    @staticmethod
    def _extract_tables(doc: _Document) -> list[ParsedTable]:
        tables: list[ParsedTable] = []
        for idx, t in enumerate(doc.tables):
            rows = [[cell.text for cell in row.cells] for row in t.rows]
            tables.append(
                ParsedTable(
                    rows=rows,
                    page_no=None,  # DOCX has no native page concept
                    order_in_page=idx,
                )
            )
        return tables

    @staticmethod
    def _extract_images(docx_path: Path) -> list[ExtractedImage]:
        """Read ``word/media/*`` straight from the DOCX zip.

        Each image is hashed (sha256) and written to a uniquely-named temp file;
        the caller is responsible for cleanup if persistence is desired.
        """
        from PIL import Image

        images: list[ExtractedImage] = []
        try:
            with zipfile.ZipFile(docx_path) as zf:
                media_names = sorted(
                    n for n in zf.namelist() if n.startswith("word/media/") and not n.endswith("/")
                )
                for idx, name in enumerate(media_names):
                    data = zf.read(name)
                    sha = hashlib.sha256(data).hexdigest()

                    width, height, mime = _inspect_image(data)

                    suffix = Path(name).suffix or ".bin"
                    with tempfile.NamedTemporaryFile(
                        prefix=f"kdp_docx_{idx}_",
                        suffix=suffix,
                        delete=False,
                    ) as fh:
                        fh.write(data)
                    tmp_path_str = fh.name

                    images.append(
                        ExtractedImage(
                            page_no=None,
                            section_no=None,
                            bbox=None,
                            bbox_unit="none",
                            order_in_page=idx,
                            text_before="",
                            text_after="",
                            section_title=None,
                            file_path=tmp_path_str,
                            sha256=sha,
                            width=width,
                            height=height,
                            size_bytes=len(data),
                            mime_type=mime,
                            detected_caption=None,
                            caption_method=None,
                            caption_pattern_score=0.0,
                        )
                    )
        except zipfile.BadZipFile:
            # Corrupt zip — surface as ParseError from the caller's try block.
            raise

        # Suppress unused-import warning when type-checking-only
        _ = Image
        return images

    @staticmethod
    def _extract_metadata(doc: _Document, path: Path) -> ParseMetadata:
        props = doc.core_properties
        return ParseMetadata(
            format="docx",
            file_path=path,
            title=props.title or None,
            author=props.author or None,
            created_at=props.created,
            modified_at=props.modified,
        )


def _inspect_image(data: bytes) -> tuple[int, int, str]:
    """Return ``(width, height, mime_type)`` for raw image bytes.

    Falls back to ``(0, 0, "application/octet-stream")`` on any error so a
    pathological media entry doesn't blow up the whole parse.
    """
    from PIL import Image

    try:
        with Image.open(BytesIO(data)) as img:
            width, height = img.size
            mime = Image.MIME.get(img.format or "", "application/octet-stream")
            return width, height, mime
    except Exception:
        return 0, 0, "application/octet-stream"


# ─────────────────────────────────────────────────────────────────────────────
# Auto-register on import
# ─────────────────────────────────────────────────────────────────────────────

register_parser(DocxParser())
